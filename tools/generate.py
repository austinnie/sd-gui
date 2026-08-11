#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
通用生成器：根据提示词库自动出图（支持分层+扁平，全部生成）
用法：python generate.py <风格名称>
例如：python generate.py pure_serene

指定生成数量：
  python generate.py pure_serene_v2 -n 10
  python generate.py pure_serene_v2 --count 20

生成模式：
  --img2img, --i2i    图生图模式（默认，需要 input.jpg）
  --txt2img, --t2i    文生图模式（无需参考图，从零生成）

步数控制：
  --steps <数字>      临时指定生成步数（优先于 config.py 中的 STEPS）

图生图输入：
  --input <文件路径>   指定图生图的参考图（不指定则默认读取 config.py 中的 input.jpg）
"""
import os
import sys
import io

# ========== 修复 Windows 终端编码问题 ==========
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
# =================================================

import time
import cv2
import numpy as np
import torch
import time
import random
from PIL import Image



from datetime import datetime
from diffusers import StableDiffusionPipeline, EulerDiscreteScheduler

# 确保 tools 目录在路径中
CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
if CURRENT_DIR not in sys.path:
    sys.path.insert(0, CURRENT_DIR)

# ✅ 添加项目根目录到路径（让 utils 可以被导入）
PROJECT_ROOT = os.path.dirname(CURRENT_DIR)
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

# ✅ 验证路径
print(f"📁 PROJECT_ROOT: {PROJECT_ROOT}")

# ✅ 引入你项目内置的 BLIP 后端
#from gui.tabs.interrogate.backends.blip import BlipBackend

# ✅ 导入 utils
from utils.imagemeta_cleaner import smart_clean_image
from utils.exif_injector import inject_exif
from utils.photo_realistic import make_photo_realistic


# ========== ✅ 导入配置 ==========
from tools.config import (
    SD_MODEL_PATH, STEPS, MAX_LIMIT, INPUT_IMAGE_NAME, DEFAULT_STRENGTH,
    REMOVE_AI_TRACES, AI_CLEAR_METADATA, AI_INJECT_EXIF, AI_REALISTIC,
    AI_CAMERA, AI_STRENGTH, AI_STYLE, AI_RANDOMIZE,
    AI_FINGERPRINT_OBFUSCATION, AI_DISTORTION_STRENGTH,
    AI_CHROMATIC_ABERRATION, AI_CHROMATIC_STRENGTH,
    AI_REALISTIC_NOISE, AI_NOISE_ISO_BASE, AI_NOISE_RANDOMIZE,
    AI_MINOR_CROP, AI_CROP_PERCENT,
    AUTO_DETECT_STYLE, SKETCH_KEYWORDS,
    # ========== 🆕 导入互斥开关（去掉 0~3 死路径导入） ==========
    USE_OPENVINO_MODEL, ACTIVE_MODEL,
    # 🛑 注意：不要在这里导入 SD_OV_MODEL_PATH, SD_MODEL_PATH_0/1/2/3 
    AI_APPRECIATION_ENGINE
)

from tools.config import SCHEDULER_TYPE

print(f"📊 STEPS = {STEPS}")
print(f"📷 AI_CAMERA = {AI_CAMERA}")

from prompts_config import STYLE_PROMPTS

# ==================== ⚙️ 安全开关 ====================
SAFE_MODE = True  
# 安全模式策略：
#   "simple" = 简单模式：在提示词后加 "wearing clothes"
#   "filter" = 过滤模式：移除露骨词汇 (nude, naked, explicit, pornographic, sex, hentai)
SAFE_MODE_STRATEGY = "filter"  # 可选: "simple" 或 "filter"

# 是否启用去水印
REMOVE_WATERMARK = True

# ==================== ⚙️ 内容文本开关 ====================
# 是否启用 content_texts 字段（将文本内容添加到提示词中）
USE_CONTENT_TEXTS = True  # 默认关闭，设为 True 开启
# ========================================================

# ==================== 🛠️ 工具函数 ====================

def print_usage():
    """打印使用说明"""
    print("\n" + "="*60)
    print("📖 通用生成器使用说明")
    print("="*60)
    print("\n用法：")
    print("  python generate.py <风格名称>")
    print("  python generate.py <风格名称> -n <数量>")
    print("  python generate.py <风格名称> --count <数量>")
    print("  python generate.py <风格名称> --steps <数字>")
    print("  python generate.py <风格名称> --input <参考图路径>")
    print("\n生成模式：")
    print("  --img2img, --i2i    图生图模式（默认，需要 input.jpg）")
    print("  --txt2img, --t2i    文生图模式（无需参考图，从零生成）")
    print("\n步数控制：")
    print("  --steps <数字>      指定生成步数（不指定则使用 config.py 中的 STEPS）")
    print("\n图生图输入：")
    print("  --input <文件路径>  指定参考图（不指定则使用 config.py 中的默认 input.jpg）")
    print("\n提示词库选择：")
    print("  --use-old           使用旧版提示词库 (prompts/ 目录)")
    print("  （默认使用新版安全提示词库 prompts_new/ 目录）")
    print("\n示例：")
    print("  python generate.py anime_xxx_v3 --steps 30 -n 5   # 30步生成5张")
    print("  python generate.py anime_xxx_v3 --txt2img         # 使用config默认步数")
    print("  python generate.py anime_xxx_v3 --img2img --input my_pic.png -n 3 # 用指定图生图")
    print("  python generate.py --use-old babata_poses -n 3    # 使用旧版提示词库")
    print("\n其他命令：")
    print("  python generate.py --list     显示所有可用风格（分屏）")
    print("  python generate.py -l         显示所有可用风格（分屏）")
    print("  python generate.py --search <关键词>  搜索风格")
    print("  python generate.py -s <关键词>       搜索风格")
    print("\n💡 提示：")
    print("  - 图生图模式：需要 input.jpg 作为参考图")
    print("  - 文生图模式：不需要参考图，完全根据提示词生成")
    print("  - 文生图模式会自动随机选择尺寸，增加多样性")
    print("  - 去水印功能: " + ("✅ 已开启" if REMOVE_WATERMARK else "❌ 已关闭"))
    print("="*60)

def print_style_list():
    """分屏显示风格列表"""
    styles = list(STYLE_PROMPTS.keys())
    total = len(styles)
    
    print("\n" + "="*60)
    print(f"📋 当前支持的风格列表（共 {total} 个）")
    print("="*60)
    
    # 每页显示数量
    page_size = 20
    total_pages = (total + page_size - 1) // page_size
    
    for page in range(total_pages):
        start = page * page_size
        end = min(start + page_size, total)
        
        print(f"\n📄 第 {page+1}/{total_pages} 页")
        print("-"*40)
        
        for i in range(start, end):
            style_name = styles[i]
            folder_info = ""
            if style_name in STYLE_PROMPTS:
                folder = STYLE_PROMPTS[style_name].get("folder", "")
                if folder:
                    folder_info = f" -> {folder}"
            print(f"  {i+1:3d}. {style_name}{folder_info}")
        
        print("-"*40)
        print(f"显示 {end-start} 个，共 {total} 个风格")
        
        if page < total_pages - 1:
            input("\n按 Enter 继续查看下一页...")
    
    print("\n" + "="*60)
    print("💡 使用方式：python generate.py <风格名称> [-n <数量>]")
    print("="*60)

def find_input_image(custom_input=None):
    """查找参考图，支持自定义路径或默认路径"""
    if custom_input:
        if os.path.exists(custom_input):
            return custom_input
        else:
            print(f"⚠️ 警告：找不到指定的输入文件 '{custom_input}'，尝试查找默认 input 文件...")
    
    # 默认查找 INPUT_IMAGE_NAME
    for ext in [".jpg", ".jpeg", ".png", ".webp", ".bmp"]:
        path = os.path.join(CURRENT_DIR, INPUT_IMAGE_NAME + ext)
        if os.path.exists(path):
            return path
    return None

# ==================== 去水印功能 ====================
def remove_watermark(image_path):
    """
    检测并去除图片水印
    返回: PIL Image 对象
    """
    if not REMOVE_WATERMARK:
        print("[系统] 去水印功能已关闭，直接使用原图")
        return Image.open(image_path).convert('RGB')
    
    print("\n[AI预处理] 检测并去除图片水印...")
    
    # ✅ 修复中文路径：使用 imdecode 代替 imread
    try:
        # 方法1：用 imdecode 读取
        with open(image_path, 'rb') as f:
            img_bytes = np.frombuffer(f.read(), np.uint8)
        img = cv2.imdecode(img_bytes, cv2.IMREAD_COLOR)
        
        if img is None:
            # 方法2：降级使用 PIL
            pil_img = Image.open(image_path).convert('RGB')
            img = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
            if img is None:
                raise ValueError("无法读取图片")
    except Exception as e:
        print(f"⚠️ 读取图片失败，跳过水印检测: {e}")
        return Image.open(image_path).convert('RGB')
    
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    
    # 检测白色/亮色区域（常见水印特征）
    _, mask = cv2.threshold(gray, 230, 255, cv2.THRESH_BINARY)
    kernel = np.ones((5,5), np.uint8)
    mask = cv2.dilate(mask, kernel, iterations=1)
    
    # 计算白色区域占比
    white_pixel_ratio = np.sum(mask > 0) / mask.size
    
    # 如果白色区域过少或过多，认为没有明显水印
    if white_pixel_ratio < 0.01 or white_pixel_ratio > 0.2:
        print("✅ 未检测到明显水印，继续生成。")
        return Image.open(image_path).convert('RGB')

    print("⚠️ 检测到水印，正在使用 OpenCV 修复去除...")
    result = cv2.inpaint(img, mask, 3, cv2.INPAINT_TELEA)
    print("✅ 水印去除完成！")
    return Image.fromarray(cv2.cvtColor(result, cv2.COLOR_BGR2RGB))
    
# ==================== 🚀 核心：加载模型管道 ====================
# tools/generate.py
# 在 setup_pipeline() 函数中，替换普通模型分支

def setup_pipeline():
    print(f"\n[系统] 正在加载 AI 模型...")

    # ========== 🆕 统一使用 config.py 的 SD_MODEL_PATH ==========
    if USE_OPENVINO_MODEL:
        # 【开】OpenVINO 分支
        try:
            from tools.config import SD_OV_MODEL_PATH
            model_path = SD_OV_MODEL_PATH
        except ImportError:
            print("❌ 错误：试图使用 OpenVINO，但 config.py 中没有定义 SD_OV_MODEL_PATH。")
            sys.exit(1)

        print(f"⚡ [配置] 使用 OpenVINO 加速模式")
        print(f"   📂 模型路径: {model_path}")
        
        try:
            from optimum.intel import OVStableDiffusionPipeline
            print("⚡ 尝试加载 OpenVINO 接口...")
            pipe = OVStableDiffusionPipeline.from_pretrained(model_path, compile=False, export=True)
            print("✅ OpenVINO 模型加载成功！")
        except Exception as e:
            print(f"❌ OpenVINO 加载失败: {e}")
            sys.exit(1)
    else:
        # 【关】普通模型分支 - 🆕 使用新的 SD_MODEL_PATH
        print(f"⚡ [配置] 使用普通模型模式")
        
        # 🆕 从 config.py 获取最终模型路径
        try:
            from tools.config import SD_MODEL_PATH
            model_path = SD_MODEL_PATH
            print(f"   📂 模型路径: {model_path}")
        except ImportError:
            print("❌ 错误：无法从 config.py 导入 SD_MODEL_PATH")
            sys.exit(1)
        
        # 如果是目录，查找 .safetensors 文件
        if os.path.isdir(model_path):
            import glob
            safetensors_files = glob.glob(os.path.join(model_path, "*.safetensors"))
            if safetensors_files:
                model_path = safetensors_files[0]
                print(f"   🔍 自动定位到目录内的模型文件: {os.path.basename(model_path)}")
            else:
                print(f"❌ 错误：在目录 {model_path} 中找不到任何 .safetensors 文件。")
                sys.exit(1)
                
        try:
            pipe = StableDiffusionPipeline.from_single_file(
                model_path,
                torch_dtype=torch.float32,
                safety_checker=None,
                requires_safety_checker=False,
                use_safetensors=True
            )
            pipe.to("cpu")
            print("✅ 普通模型加载成功！")

            # ================= 🆕 加载 LoRA =================
            try:
                from tools.config import FINAL_LORA_LIST
                
                if FINAL_LORA_LIST:
                    print(f"   📦 准备加载 {len(FINAL_LORA_LIST)} 个 LoRA...")
                    
                    adapter_names = []
                    adapter_weights = []
                    
                    for i, lora_info in enumerate(FINAL_LORA_LIST):
                        lora_path = lora_info['path']
                        lora_weight = lora_info['weight']
                        
                        if os.path.exists(lora_path):
                            adapter_name = f"lora_{i}"
                            print(f"      🔗 加载 LoRA {i+1}: {os.path.basename(lora_path)} (权重: {lora_weight})")
                            
                            pipe.load_lora_weights(lora_path, adapter_name=adapter_name)
                            
                            adapter_names.append(adapter_name)
                            adapter_weights.append(lora_weight)
                    
                    if adapter_names:
                        pipe.set_adapters(adapter_names, adapter_weights=adapter_weights)
                        print(f"      ✅ 全部 {len(adapter_names)} 个 LoRA 加载成功！")
                    else:
                        print("   ⚠️ 没有找到任何有效的 LoRA 文件。")
                else:
                    print("   ℹ️ 未配置 LoRA，将使用裸模型出图")
                    
            except ImportError:
                print("   ℹ️ config.py 未定义 LoRA 配置，跳过 LoRA 加载")
            # ========================================================
            
        except Exception as e:
            print(f"❌ 普通模型加载失败: {e}")
            sys.exit(1)

    # ======================= 🚀 动态加载 IP-Adapter =======================
    # 注意：IP-Adapter 必须从 pipe 加载，并在调度器设置之前加载
    # ================= 🆕 预加载 IP-Adapter (下载权重) =================
    try:
        print("   🧬 正在加载/检查 IP-Adapter 模型权重...")
        pipe.load_ip_adapter(
            "h94/IP-Adapter",
            subfolder="models",
            weight_name="ip-adapter_sd15.safetensors",
        )
        print(f"   ✅ IP-Adapter 权重加载/检查完成！")
    except Exception as e:
        print(f"   ⚠️ IP-Adapter 加载失败 (如果没有使用 --ip_adapter 可忽略此警告): {e}")
    # ====================================================================
    # ========================================================================
    
    # ======================= 🚀 动态采样器加载 =======================
    try:
        pipe.enable_vae_slicing()
        pipe.enable_attention_slicing()

        # 从 config.py 导入我们刚才加的采样器变量
        from tools.config import FINAL_SCHEDULER
        
        print(f"   🎛️ 正在加载采样器: {FINAL_SCHEDULER}...")
        
        if FINAL_SCHEDULER == "Euler":
            pipe.scheduler = EulerDiscreteScheduler.from_config(pipe.scheduler.config)
        elif FINAL_SCHEDULER == "EulerAncestral":
            from diffusers import EulerAncestralDiscreteScheduler
            pipe.scheduler = EulerAncestralDiscreteScheduler.from_config(pipe.scheduler.config)
        elif FINAL_SCHEDULER == "DPM++ 2M":
            from diffusers import DPMSolverMultistepScheduler
            pipe.scheduler = DPMSolverMultistepScheduler.from_config(pipe.scheduler.config)
        elif FINAL_SCHEDULER == "DPM++ 2M Karras":
            from diffusers import DPMSolverMultistepScheduler
            pipe.scheduler = DPMSolverMultistepScheduler.from_config(pipe.scheduler.config, use_karras_sigmas=True)
        elif FINAL_SCHEDULER == "DPM++ SDE Karras":
            from diffusers import DPMSolverSDEscheduler
            pipe.scheduler = DPMSolverSDEScheduler.from_config(pipe.scheduler.config, use_karras_sigmas=True)
        elif FINAL_SCHEDULER == "DDIM":
            from diffusers import DDIMScheduler
            pipe.scheduler = DDIMScheduler.from_config(pipe.scheduler.config)
        elif FINAL_SCHEDULER == "PNDM":
            from diffusers import PNDMScheduler
            pipe.scheduler = PNDMScheduler.from_config(pipe.scheduler.config)
        elif FINAL_SCHEDULER == "LMS":
            from diffusers import LMSDiscreteScheduler
            pipe.scheduler = LMSDiscreteScheduler.from_config(pipe.scheduler.config)
        elif FINAL_SCHEDULER == "Heun":
            from diffusers import HeunDiscreteScheduler
            pipe.scheduler = HeunDiscreteScheduler.from_config(pipe.scheduler.config)
        elif FINAL_SCHEDULER == "UniPC":
            from diffusers import UniPCMultistepScheduler
            pipe.scheduler = UniPCMultistepScheduler.from_config(pipe.scheduler.config)
        else:
            print(f"⚠️ 未知采样器 '{FINAL_SCHEDULER}'，回退到默认 Euler")
            pipe.scheduler = EulerDiscreteScheduler.from_config(pipe.scheduler.config)
            
        print(f"   ✅ 采样器加载完成！")
        
    except Exception as e:
        print(f"⚠️ 注意：采样器加载失败，使用默认配置。错误: {e}")
    # ================================================================
        
    print("[系统] 模型管道已就绪！")
    return pipe
    
def build_prompt(config):
    """
    分层构建提示词
    支持三种格式：
    1. 分层格式：subjects + styles + moods
    2. 扁平格式：只有 subjects（兼容旧配置）
    3. 内容文本扩展：subjects + styles + moods + content_texts（需开启 USE_CONTENT_TEXTS）
    """
    if "styles" in config and "moods" in config:
        subject = random.choice(config["subjects"])
        style = random.choice(config["styles"])
        mood = random.choice(config["moods"])
        
        # 如果开启且存在 content_texts，随机选一句添加
        if USE_CONTENT_TEXTS and "content_texts" in config and config["content_texts"]:
            text = random.choice(config["content_texts"])
            # ✨ 修改点：去掉死板的 "calligraphy text: "，将文字融入画面描述
            prompt = f"{subject}, {style}, {mood}, the scroll features the Chinese characters '{text}' written in flowing calligraphy"
            print(f"   📜 已添加内容文本: {text[:20]}...")
        else:
            prompt = f"{subject}, {style}, {mood}"
        return prompt, "分层"
    else:
        prompt = random.choice(config["subjects"])
        return prompt, "扁平"
        
def generate_style(pipe, init_image, prompt, output_filename, strength, mode="img2img", steps=STEPS, target_style="unknown", enable_ip_adapter=False, ip_adapter_scale=1.0):
    """
    生成单张图片
    mode: "img2img" 或 "txt2img"
    steps: 当前生成使用的步数
    """
    
    # ========== 📊 详细调试信息 ==========
    print(f"\n{'='*50}")
    print(f"📊 [调试] 参数详情:")
    print(f"  ├─ 图生图强度 (strength): {strength}")
    print(f"  ├─ 照片真实化强度 (AI_STRENGTH): {AI_STRENGTH}")
    print(f"  ├─ 紫边模拟强度 (AI_CHROMATIC_STRENGTH): {AI_CHROMATIC_STRENGTH}")
    print(f"  ├─ 消除AI痕迹 (REMOVE_AI_TRACES): {REMOVE_AI_TRACES}")
    print(f"  └─ 迭代步数 (steps): {steps}")
    print(f"{'='*50}\n")
    
    import random  # ✅ 添加这一行  
    from PIL import Image    
    max_limit = MAX_LIMIT
    
    if mode == "img2img":
        # 图生图：使用原图尺寸
        w, h = init_image.size
        if w > max_limit or h > max_limit:
            if w > h:
                scale = max_limit / w
            else:
                scale = max_limit / h
            w, h = int(w * scale), int(h * scale)
        w, h = ((w+31)//64)*64, ((h+31)//64)*64
        
        # ✅ 使用兼容写法
        try:
            image = init_image.resize((w, h), Image.Resampling.LANCZOS)
        except AttributeError:
            image = init_image.resize((w, h), Image.LANCZOS)
            
        print(f"[图生图] {os.path.basename(output_filename)} ({w}x{h})")
    else:
        # 文生图：随机选择尺寸，增加多样性
        aspect_ratios = [
            (512, 512), (512, 576), (576, 512),
            (512, 640), (640, 512), 
            (512, 768), (768, 512), 
            (576, 768), (768, 576),
            (448, 640), (640, 448),
            # --- 手机壁纸（竖屏） ---
            (768, 1360), (1080, 1920), (768, 1024),
            # --- 电脑壁纸（横屏） ---
            (1024, 576), (1280, 720), (1360, 768),
            # --- 经典正方形 ---
            (768, 768), (1024, 1024)            
        ]
        w, h = random.choice(aspect_ratios)
        # 限制最大尺寸
        if w > max_limit: w = max_limit
        if h > max_limit: h = max_limit
        w, h = ((w+31)//64)*64, ((h+31)//64)*64
        image = None
        print(f"[文生图] {os.path.basename(output_filename)} ({w}x{h})")

    # ========== 安全模式处理 ==========
    if SAFE_MODE:
        if SAFE_MODE_STRATEGY == "simple":
            # 简单模式：直接加 wearing clothes
            full_prompt = f"{prompt}, wearing clothes"
            print(f"🛡️ [安全模式] 策略: 简单 (附加 wearing clothes)")
        elif SAFE_MODE_STRATEGY == "filter":
            # 过滤模式：移除露骨词汇
            safe_prompt = prompt.replace("nude", "").replace("naked", "").replace("explicit", "").replace("pornographic", "")
            safe_prompt = ", ".join([p for p in safe_prompt.split(",") if "sex" not in p.lower() and "hentai" not in p.lower() and "penetration" not in p.lower()])
            # 清理多余的逗号和空格
            safe_prompt = ", ".join([p.strip() for p in safe_prompt.split(",") if p.strip()])
            full_prompt = safe_prompt if safe_prompt.strip() else prompt
            print(f"🛡️ [安全模式] 策略: 过滤 (移除露骨词汇)")
        else:
            # 默认：简单模式
            full_prompt = f"{prompt}, wearing clothes"
            print(f"🛡️ [安全模式] 策略: 默认 (附加 wearing clothes)")
        
        # 统一使用强化版负面提示词
        neg_prompt = (
            "worst quality, low quality, ugly, deformed, blurry, watermark, signature, logo, brand, "
            "bad hands, extra fingers, missing fingers, fused fingers, deformed hands, "
            "mutated hands, poorly drawn hands, six fingers, eleven fingers, "
            "bad anatomy, malformed limbs, extra limbs, missing limbs, "
            "bad proportions, disfigured, gross proportions, "
            "bad feet, extra toes, missing toes, fused toes, "
            "jumbled text, gibberish characters, messy ink, smudged writing, illegible scribbles, unreadable signs"            
        )
    else:
        full_prompt = prompt
        neg_prompt = (
            "worst quality, low quality, ugly, deformed, blurry, watermark, signature, logo, brand, "
            "bad hands, extra fingers, missing fingers, fused fingers, deformed hands, "
            "mutated hands, poorly drawn hands, six fingers, eleven fingers, "
            "bad anatomy, malformed limbs, extra limbs, missing limbs, "
            "bad proportions, disfigured, gross proportions, "
            "bad feet, extra toes, missing toes, fused toes, "
            "extra arm, extra hand, missing arm, missing hand, "
            "fused hand, extra digit, wrong finger count, "
            "deformed finger, twisted finger, broken finger, "
            "claw hand, abnormal hand, mutant hand, "
            "bad pose, unnatural pose, contorted body, twisted body, "
            "jumbled text, gibberish characters, messy ink, smudged writing, illegible scribbles, meaningless strokes"
        )
        print(f"🔓 [自由模式已启用]")
    
    # ========== 🧠 自动添加解剖约束 ==========
    # 检测是否包含复杂姿势关键词，自动添加约束
    complex_pose_keywords = [
        "sex", "posing", "bending", "kneeling", "lying", 
        "standing", "spooning", "riding", "missionary", 
        "doggy", "cowgirl", "spread", "bent over", 
        "on top", "from behind", "oral", "blowjob", 
        "cunnilingus", "group", "threesome"
    ]
    
    # 检测是否包含多人关键词
    multi_person_keywords = ["two", "multiple", "group", "couple", "pair", "man and woman", "both bodies"]
    
    # 检测是否包含天使/翅膀关键词
    wing_keywords = ["angel wings", "feathered wings", "bird wings", "butterfly wings", "dragon wings", "wings spread", "winged figure"]
    
    # ✨ 草稿阻断逻辑：遇到结构草稿，清空翅膀检测
    sketch_keywords = ["sketch", "pencil", "draft", "wireframe", "construction", "anatomy", "lineart", "structural"]
    if any(keyword in prompt.lower() for keyword in sketch_keywords):
        wing_keywords = []
        
    # 构建约束
    constraints = []
    
    # 复杂姿势约束
    if any(keyword in prompt.lower() for keyword in complex_pose_keywords):
        constraints.append("natural body position")
        constraints.append("correct anatomy")
        constraints.append("realistic hands and feet")
    
    # 多人场景约束
    if any(keyword in prompt.lower() for keyword in multi_person_keywords):
        constraints.append("two hands per person")
        constraints.append("two feet per person")
        constraints.append("normal proportions")
    
    # 翅膀约束
    if any(keyword in prompt.lower() for keyword in wing_keywords):
        constraints.append("symmetrical wings")
        constraints.append("beautiful feathered wings")
    
    # 如果有约束，添加到提示词末尾
    if constraints:
        constraint_text = ", ".join(constraints)
        full_prompt = f"{full_prompt}, {constraint_text}"
        print(f"   🧠 已添加解剖约束: {constraint_text}")
    
    print(f"[{datetime.now().strftime('%H:%M:%S')}] 提示词: {full_prompt[:80]}...")
    print(f"  步数: {steps}")
    
    generator = torch.Generator("cpu").manual_seed(int(time.time_ns() % 1000000000))
    
    # ========== 🚀 统一使用 inference_kwargs + IP-Adapter 支持 ==========
    inference_kwargs = {
        "prompt": full_prompt,
        "negative_prompt": neg_prompt,
        "num_inference_steps": steps,
        "guidance_scale": 7.5,
        "generator": generator,
        "width": w,
        "height": h,
    }

    if mode == "img2img":
        inference_kwargs["image"] = image
        inference_kwargs["strength"] = strength
        
        # 🆕 如果启用了 IP-Adapter 并且有参考图，则把图片作为参考传进去
        if enable_ip_adapter and image is not None:
            # 注意：ip_adapter_image 需要传入 PIL Image 对象
            inference_kwargs["ip_adapter_image"] = image
            print(f"   🧬 [IP-Adapter] 已注入参考图 (Scale: {ip_adapter_scale})")

    # 执行统一的推理
    result = pipe(**inference_kwargs)
    # =====================================================================
    
    # 保存图片
    result.images[0].save(output_filename, quality=95)

    # ========== 🆕 检测风格 ==========
    prompt_lower = prompt.lower()
    is_sketch = False
    if AUTO_DETECT_STYLE:
        is_sketch = any(kw in prompt_lower for kw in SKETCH_KEYWORDS)
        # 也检查目标风格名称
        if not is_sketch:
            is_sketch = any(kw in target_style.lower() for kw in SKETCH_KEYWORDS)
    
    if is_sketch:
        print(f"\n🎨 检测到素描/线稿风格，仅清除元数据，跳过相机相关处理")
    # ================================
    
    # ========== 🆕 消除AI痕迹 - 后期处理 ==========
    if REMOVE_AI_TRACES:
        try:
            print(f"\n📷 消除AI痕迹处理...")
            final_path = output_filename
            
            # 1️⃣ 清除元数据（如果需要）
            if AI_CLEAR_METADATA:

                # 转换为JPG并清除元数据
                jpg_path = output_filename.replace('.png', '.jpg')
                final_path = smart_clean_image(
                    final_path, 
                    output_path=jpg_path,
                    method='jpg',
                    jpg_quality=92
                )
                print(f"   ✅ 元数据已清除 -> JPG")

            # 2️⃣ 照片真实化（添加噪点、暗角、锐化）- 由 photo_realistic 单独处理
            if AI_REALISTIC and not is_sketch:
                
                final_path = make_photo_realistic(
                    final_path,
                    final_path,
                    camera=AI_CAMERA,
                    style="portrait",
                    inject_exif_data=False,
                    randomize=True,
                    strength=AI_STRENGTH,
                    add_noise_flag=AI_REALISTIC_NOISE  # 👈 把 config.py 里的配置真实地传进去
                )
                print(f"   ✅ 照片真实化完成 (强度: {AI_STRENGTH})")
            
            # 3️⃣ 注入 EXIF（如果开启了，且之前没有注入）
            if AI_INJECT_EXIF and not is_sketch:
                time.sleep(0.5) # 等待 0.5 秒避免文件锁
                try:
                    final_path = inject_exif(
                        final_path,
                        final_path,
                        camera=AI_CAMERA,
                        style="portrait",
                        randomize=True
                    )
                    print(f"   ✅ EXIF 已注入")
                except Exception as e:
                    print(f"   ⚠️ EXIF 注入过程抛异常，已跳过: {e}")

            # ====== 下面是完全独立的 3 个图像处理滤镜 ======
            # 它们不再被捆绑在“指纹混淆”下，各自独立运作

            # 4️⃣ 紫边模拟（独立）
            if AI_CHROMATIC_ABERRATION and not is_sketch:
                try:
                    print(f"   🔬 模拟紫边...")
                    from PIL import Image
                    import numpy as np
                    
                    img = Image.open(final_path)
                    # 转为numpy数组处理
                    arr = np.array(img).astype(np.float32)
                    h, w = arr.shape[:2]
                    strength = AI_CHROMATIC_STRENGTH
                    
                    # 在图像边缘添加红/蓝通道偏移（紫边特征）
                    for y in range(h):
                        for x in range(w):
                            # 计算距离边缘的距离
                            dist_from_edge = min(x, w-1-x, y, h-1-y)
                            if dist_from_edge < 40:
                                # 越靠近边缘，紫边越明显
                                shift_factor = (40 - dist_from_edge) / 40
                                shift = shift_factor * strength * random.uniform(0.5, 1.0)
                                # 红色通道偏移（紫色倾向）
                                arr[y, x, 0] += random.uniform(-shift, shift * 0.5)  # R
                                arr[y, x, 2] += random.uniform(-shift * 0.5, shift)  # B
                    
                    # 裁剪到有效范围
                    arr = np.clip(arr, 0, 255).astype(np.uint8)
                    
                    # 确保维度是 HWC 且转为 RGB 标准格式
                    if arr.ndim == 3 and arr.shape[2] == 3:
                        img = Image.fromarray(arr).convert('RGB')
                    else:
                        img = Image.fromarray(arr[:, :, :3]).convert('RGB')
                    
                    img.save(final_path, quality=92)
                    print(f"      ✅ 紫边模拟完成 (强度: {strength})")
                    
                except Exception as e:
                    print(f"   ⚠️ 紫边模拟失败: {e}")

            # 5️⃣ 真实噪点（独立，不会因为紫边失败而受影响）
            if AI_REALISTIC_NOISE and not is_sketch:
                try:
                    print(f"   📸 添加真实噪点...")
                    import cv2
                    import numpy as np
                    from PIL import Image
                    
                    # 确定ISO值
                    if AI_NOISE_RANDOMIZE:
                        iso = AI_NOISE_ISO_BASE + random.randint(-200, 200)
                        iso = max(100, min(1600, iso))
                    else:
                        iso = AI_NOISE_ISO_BASE
                    
                    img = Image.open(final_path)
                    # 强制转换为 uint8，绝对防止 OpenCV 报错
                    img_np = np.array(img).astype(np.uint8)
                    # 如果是 4 通道（RGBA），扔掉透明通道
                    if img_np.ndim == 3 and img_np.shape[2] == 4:
                        img_np = img_np[:, :, :3]
                    
                    # 转换为OpenCV格式
                    img_cv = cv2.cvtColor(img_np, cv2.COLOR_RGB2BGR)
                    
                    # 基于ISO的噪声强度
                    noise_std = 0.005 * (iso / 100) ** 0.5
                    
                    # 高斯噪声（模拟传感器热噪声）
                    gaussian_noise = np.random.normal(0, noise_std * 255, img_cv.shape)
                    
                    # 散粒噪声（泊松分布模拟，光子噪声）
                    shot_noise = np.random.poisson(np.abs(img_cv) * 0.005) * 0.1
                    
                    # 合并噪声
                    img_cv = img_cv + gaussian_noise + shot_noise
                    
                    # 暗部噪点增强
                    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                    dark_mask = gray < 80
                    if np.any(dark_mask):
                        dark_noise = np.random.normal(0, noise_std * 255 * 0.5, img_cv.shape)
                        img_cv[dark_mask] = img_cv[dark_mask] + dark_noise[dark_mask]
                    
                    # 裁剪到有效范围
                    img_cv = np.clip(img_cv, 0, 255).astype(np.uint8)
                    img = Image.fromarray(cv2.cvtColor(img_cv, cv2.COLOR_BGR2RGB))
                    img.save(final_path, quality=92)
                    print(f"      ✅ 真实噪点添加完成 (ISO: {iso})")
                    
                except Exception as e:
                    print(f"   ⚠️ 真实噪点添加失败: {e}")

            # 6️⃣ 轻微裁剪（独立）
            if AI_MINOR_CROP:
                try:
                    print(f"   ✂️ 轻微裁剪...")
                    from PIL import Image
                    import random
                    
                    img = Image.open(final_path)
                    w, h = img.size
                    
                    crop_pct = AI_CROP_PERCENT * random.uniform(0.5, 1.5)
                    crop_w = int(w * crop_pct)
                    crop_h = int(h * crop_pct)
                    
                    crop_w = max(5, min(crop_w, int(w * 0.05)))
                    crop_h = max(5, min(crop_h, int(h * 0.05)))
                    
                    # 随机选择裁剪位置
                    corners = [
                        (0, 0), (0, crop_h), (crop_w, 0), (crop_w, crop_h),
                    ]
                    if random.random() < 0.5:
                        left = random.randint(0, crop_w)
                        top = random.randint(0, crop_h)
                    else:
                        left, top = random.choice(corners)
                    
                    right = w - random.randint(0, crop_w)
                    bottom = h - random.randint(0, crop_h)
                    
                    if right > left + 50 and bottom > top + 50:
                        img = img.crop((left, top, right, bottom))
                        img = img.resize((w, h), Image.Resampling.LANCZOS)
                        img.save(final_path, quality=92)
                        print(f"      ✅ 轻微裁剪完成 (裁切: {crop_pct*100:.1f}%, 位置: {left},{top})")
                    else:
                        print(f"      ⚠️ 裁剪跳过 (区域无效)")
                        
                except Exception as e:
                    print(f"   ⚠️ 轻微裁剪失败: {e}")
            
            # 如果最终路径改变了，更新文件名
            if final_path != output_filename:
                # 删除原始PNG（如果存在）
                if os.path.exists(output_filename) and output_filename != final_path:
                    try:
                        os.remove(output_filename)
                    except:
                        pass
                output_filename = final_path

                
        except Exception as e:
            print(f"   ⚠️ 消除AI痕迹整体流程失败: {e}")
    # ================================================
    

    # 只要生成图片成功，就在同一目录下生成一个同名的 .txt 说明文件
    metadata_filename = output_filename.replace('.png', '.txt').replace('.jpg', '.txt')
    try:
        with open(metadata_filename, "w", encoding="utf-8") as f:
            f.write(f"【生成时间】: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"【风格名称】: {target_style}\n")
            f.write(f"【生成模式】: {mode}\n")
            f.write(f"【迭代步数】: {steps}\n")
            
            # ========== 📊 详细记录所有 strength 参数 ==========
            f.write(f"\n【📊 Strength 参数详情】:\n")
            # 🛡️ 修复：如果使用了默认强度，txt 里就如实记录默认值；如果是自定义的，就记录自定义值
            import tools.config as cfg
            if strength == 0.35 and not any('--strength' in arg or '-strength' in arg for arg in sys.argv):
                f.write(f"  ├─ 图生图强度 (img2img strength): {cfg.DEFAULT_STRENGTH} (默认值)\n")
            else:
                f.write(f"  ├─ 图生图强度 (img2img strength): {strength}\n")
            f.write(f"  ├─ 照片真实化强度 (AI_STRENGTH): {AI_STRENGTH}\n")
            f.write(f"  ├─ 紫边模拟强度 (AI_CHROMATIC_STRENGTH): {AI_CHROMATIC_STRENGTH}\n")
            f.write(f"  └─ 消除AI痕迹总开关: {'✅ 已启用' if REMOVE_AI_TRACES else '❌ 已禁用'}\n")
            
            f.write(f"\n【📝 完整正向提示词】: \n{full_prompt}\n")
            
            if mode == "img2img":
                f.write(f"\n【🖼️ 参考图路径】: {input_path if 'input_path' in locals() else '默认 input.jpg'}\n")
            
            if REMOVE_AI_TRACES:
                f.write(f"\n【🔧 消除AI痕迹配置】:\n")
                f.write(f"   - 总开关: ✅ 已启用\n")
                f.write(f"   - 元数据清理: {'✅' if AI_CLEAR_METADATA else '❌'}\n")
                
                if is_sketch:
                    f.write(f"   - 风格检测: 素描/线稿 (⚠️ 跳过相机相关处理)\n")
                    f.write(f"   - 照片真实化: ⏭️ 已跳过 (素描风格)\n")
                    f.write(f"   - EXIF注入: ⏭️ 已跳过 (素描风格)\n")
                    f.write(f"   - 紫边模拟: ⏭️ 已跳过 (素描风格)\n")
                    f.write(f"   - 真实噪点: ⏭️ 已跳过 (素描风格)\n")
                    f.write(f"   - 轻微裁剪: {'✅' if AI_MINOR_CROP else '❌'} ({AI_CROP_PERCENT*100:.1f}%)\n")
                    f.write(f"   - 指纹混淆: {'✅' if AI_FINGERPRINT_OBFUSCATION else '❌'}\n")
                else:
                    f.write(f"   - 照片真实化: {'✅' if AI_REALISTIC else '❌'}\n")
                    f.write(f"   - 相机型号: {AI_CAMERA}\n")
                    f.write(f"   - 真实化强度: {AI_STRENGTH}\n")
                    f.write(f"   - EXIF注入: {'✅' if AI_INJECT_EXIF else '❌'}\n")
                    f.write(f"   - 紫边模拟: {'✅' if AI_CHROMATIC_ABERRATION else '❌'} (强度: {AI_CHROMATIC_STRENGTH})\n")
                    f.write(f"   - 真实噪点: {'✅' if AI_REALISTIC_NOISE else '❌'} (ISO: {AI_NOISE_ISO_BASE})\n")
                    f.write(f"   - 轻微裁剪: {'✅' if AI_MINOR_CROP else '❌'} ({AI_CROP_PERCENT*100:.1f}%)\n")
                    f.write(f"   - 指纹混淆: {'✅' if AI_FINGERPRINT_OBFUSCATION else '❌'}\n")
        
        print(f"   📝 已生成提示词记录: {os.path.basename(metadata_filename)}")
    except Exception as e:
        print(f"   ⚠️ 提示词记录文件写入失败: {e}")
        
    # 在函数最后，返回最终的文件路径
    return output_filename  # ✅ 添加这行    
 
    
# ==================== 🚀 主入口 ====================

def parse_arguments(args):
    """
    解析命令行参数
    返回: (target_style, count, mode, search_keyword, steps, input_path)
    """
    target_style = None
    count = None
    mode = "img2img"
    search_keyword = None
    steps = None
    input_path = None

    # 新增参数
    clean_ai = True  # 默认启用
    no_clean = False

    enable_ip_adapter = False  # 🆕 新增
    ip_adapter_scale = 1.0     # 🆕 新增
    
    i = 1
    while i < len(args):
        arg = args[i]
        if arg in ["-n", "--count"]:
            if i + 1 < len(args):
                try:
                    count = int(args[i + 1])
                    if count <= 0:
                        print(f"❌ 数量必须大于0，当前: {count}")
                        sys.exit(1)
                    i += 2
                except ValueError:
                    print(f"❌ 无效的数字: {args[i + 1]}")
                    sys.exit(1)
            else:
                print(f"❌ 参数 {arg} 需要指定数量")
                sys.exit(1)
        elif arg in ["--txt2img", "--t2i"]:
            mode = "txt2img"
            i += 1
        elif arg in ["--img2img", "--i2i"]:
            mode = "img2img"
            i += 1
        elif arg in ["--steps"]:
            if i + 1 < len(args):
                try:
                    steps = int(args[i + 1])
                    if steps <= 0:
                        print(f"❌ 步数必须大于0，当前: {steps}")
                        sys.exit(1)
                    i += 2
                except ValueError:
                    print(f"❌ 无效的步数: {args[i + 1]}")
                    sys.exit(1)
            else:
                print(f"❌ 参数 {arg} 需要指定步数")
                sys.exit(1)
        elif arg in ["--input"]:
            if i + 1 < len(args):
                input_path = args[i + 1]
                i += 2
            else:
                print(f"❌ 参数 {arg} 需要指定文件路径")
                sys.exit(1)
        # =============== 🆕 新增 IP-Adapter 参数 ===============
        elif arg in ["--ip_adapter"]:
            enable_ip_adapter = True
            i += 1
        elif arg in ["--ip_adapter_scale"]:
            if i + 1 < len(args):
                try:
                    ip_adapter_scale = float(args[i + 1])
                    i += 2
                except ValueError:
                    print(f"❌ 无效的 IP-Adapter 权重数字: {args[i + 1]}")
                    sys.exit(1)
            else:
                print(f"❌ 参数 {arg} 需要指定数值")
                sys.exit(1)
        # =======================================================                
        elif arg in ["--search", "-s"]:
            if i + 1 < len(args):
                search_keyword = args[i + 1]
                i += 2
            else:
                print(f"❌ 参数 {arg} 需要指定搜索关键词")
                sys.exit(1)
        elif arg in ["--no-clean", "--noclean"]:
            no_clean = True
            i += 1
        elif arg in ["--use-old", "--use_old"]:
            # 这个参数在 prompts_config.py 中处理，这里直接跳过
            i += 1
        else:
            target_style = arg
            i += 1
    
    return target_style, count, mode, search_keyword, steps, input_path, no_clean, enable_ip_adapter, ip_adapter_scale
    
def main():
    # ========== 处理无参数情况 ==========
    if len(sys.argv) < 2:
        print_usage()
        sys.exit(0)
    
    # ========== 解析参数 ==========
    target_style, user_count, mode, search_keyword, user_steps, user_input, no_clean, enable_ip_adapter, ip_adapter_scale = parse_arguments(sys.argv)

    # ========== 🆕 根据命令行参数控制消除AI痕迹 ==========
    global REMOVE_AI_TRACES
    if no_clean:
        REMOVE_AI_TRACES = False
        print(f"ℹ️ 已禁用消除AI痕迹 (--no-clean)")
    else:
        print(f"ℹ️ 消除AI痕迹已启用 (相机: {AI_CAMERA}, 强度: {AI_STRENGTH})")
    # =====================================================
    
    # ========== 处理搜索 ==========
    if search_keyword:
        print(f"\n🔍 搜索包含 '{search_keyword}' 的风格：")
        print("="*60)
        found = []
        for name, config in STYLE_PROMPTS.items():
            folder = config.get("folder", "")
            if search_keyword.lower() in name.lower() or search_keyword.lower() in folder.lower():
                found.append((name, folder))
        
        if found:
            print(f"找到 {len(found)} 个匹配的风格：\n")
            for i, (name, folder) in enumerate(found, 1):
                print(f"  {i:3d}. {name} -> {folder}")
        else:
            print(f"  ❌ 没有找到包含 '{search_keyword}' 的风格")
        
        print("\n" + "="*60)
        print("💡 使用方式：python generate.py <风格名称> [-n <数量>]")
        sys.exit(0)
    
    # ========== 处理 --list 或 -l ==========
    if target_style == "--list" or target_style == "-l":
        print_style_list()
        sys.exit(0)
    
    # ========== 解析参数 ==========
    target_style, user_count, mode, search_keyword, user_steps, user_input, no_clean, enable_ip_adapter, ip_adapter_scale = parse_arguments(sys.argv)
    
    if target_style is None:
        print("❌ 请指定风格名称")
        print_usage()
        sys.exit(1)
    
    if target_style not in STYLE_PROMPTS:
        print(f"\n❌ 错误：找不到风格 '{target_style}'！")
        print("📋 可用风格列表：")
        styles = list(STYLE_PROMPTS.keys())
        for i, key in enumerate(styles[:10]):
            print(f"   - {key}")
        if len(styles) > 10:
            print(f"   ... 共 {len(styles)} 个风格")
        print("\n💡 使用 python generate.py --list 查看完整列表")
        sys.exit(1)

    # ========== 处理输入图片 ==========
    if mode == "img2img":
        # ✨ 优先使用命令行指定的输入文件
        input_path = find_input_image(user_input)
        if not input_path:
            print(f"\n❌ 图生图模式需要参考图！请用 --input 指定或把图片命名为 {INPUT_IMAGE_NAME}.jpg/.png 放在 tools 目录下！")
            return
        init_image = remove_watermark(input_path)
    else:
        init_image = None
        print(f"\n🎨 文生图模式：无需参考图，从零生成")

    # ========== 加载模型 ==========
    pipe = setup_pipeline()

    config = STYLE_PROMPTS[target_style]
    
    # ========== 计算生成数量 ==========
    # 检查是分层还是扁平
    if "styles" in config and "moods" in config:
        mode_type = "分层"
        total_possible = len(config["subjects"]) * len(config["styles"]) * len(config["moods"])
        default_count = len(config["subjects"])
        combo_info = f" ({len(config['subjects'])}×{len(config['styles'])}×{len(config['moods'])}={total_possible}种组合)"
    else:
        mode_type = "扁平"
        total_possible = len(config["subjects"])
        default_count = len(config["subjects"])
        combo_info = ""
    
    # 确定实际生成数量
    if user_count is not None:
        total_count = user_count
        count_source = f"用户指定"
        if user_count > total_possible:
            print(f"\n⚠️ 提示：您指定生成 {user_count} 张，但该风格最多只有 {total_possible} 种不同组合")
            print(f"   将生成 {total_possible} 张（全部组合）")
            total_count = total_possible
    else:
        total_count = default_count
        count_source = "全部提示词"
    
    print(f"\n🎯 正在生成风格: {target_style} -> {config['folder']}")
    print(f"📊 模式: {mode_type}{combo_info}")
    print(f"📊 本次共生成 {total_count} 张图片（{count_source}）")
    if mode_type == "分层":
        print(f"   ├─ 主体: {len(config['subjects'])} 种")
        print(f"   ├─ 风格: {len(config['styles'])} 种")
        print(f"   └─ 情绪: {len(config['moods'])} 种")
    if user_count and user_count > total_possible:
        print(f"   └─ ⚠️ 注: 实际生成 {total_possible} 张（全部组合）")
    elif user_count:
        print(f"   └─ 💡 注: 从 {total_possible} 种组合中随机选 {total_count} 张")


    # ========== 确定最终步数 (优先级：命令行 > Config > 采样器推荐) ==========
    from tools.config import FINAL_STEPS
    if user_steps is not None:
        actual_steps = user_steps
        print(f"⚙️ 步数: 命令行指定为 {user_steps} 步")
    else:
        # ✅ 修复：使用采样器推荐的 FINAL_STEPS，而不是硬编码的 STEPS
        actual_steps = FINAL_STEPS
        print(f"⚙️ 步数: 使用采样器智能推荐的 {FINAL_STEPS} 步")
    # =================================

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    # ✨ 使用配置中的 folder 作为文件夹名，避免重名
    folder_name = config['folder']
    output_root = os.path.join(CURRENT_DIR, "output", f"{folder_name}_{timestamp}")
    os.makedirs(output_root, exist_ok=True)

    # ========== 📁 新增：子文件夹分组逻辑 ==========
    # 每5张图片放一个子文件夹
    BATCH_SIZE = 5
    
    # 预计算需要创建多少个子文件夹
    total_batches = (total_count + BATCH_SIZE - 1) // BATCH_SIZE  # 向上取整
    
    # 预创建所有子文件夹
    subfolders = []
    for batch_idx in range(total_batches):
        subfolder_name = f"{batch_idx + 1:04d}"  # 从 0001 开始，四位数字
        subfolder_path = os.path.join(output_root, subfolder_name)
        os.makedirs(subfolder_path, exist_ok=True)
        subfolders.append(subfolder_path)
        print(f"📁 已创建子文件夹: {subfolder_name} (存放第 {batch_idx * BATCH_SIZE + 1} - {min((batch_idx + 1) * BATCH_SIZE, total_count)} 张)")
    
    print(f"\n📊 共 {total_count} 张图片，将分到 {total_batches} 个子文件夹中（每 {BATCH_SIZE} 张一组）\n")
    # ===============================================

    # ========== 生成循环 (终极稳定版) ==========
    from tqdm import tqdm
    # 🛡️ 核心修复：在循环开始前绝对初始化变量！
    batch_reviews = []  # 缓存当前子文件夹的点评

    # 提前检查 python-docx 是否安装，避免循环中报错导致程序崩溃
    DOCX_AVAILABLE = True
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("⚠️ [系统] 未安装 python-docx，将跳过 Word 文档生成。请运行: pip install python-docx")
        DOCX_AVAILABLE = False
    
    for i in tqdm(range(total_count), desc="生成进度"):
        prompt, prompt_mode = build_prompt(config)
        

        # 🆕 计算当前图片属于哪个子文件夹（从 0 开始计数）
        batch_index = i // BATCH_SIZE
        current_subfolder = subfolders[batch_index]

        # 🛡️ 【终极绝对防御】：只要当前路径不在我们正常的 output 目录下，立刻重定向！
        base_output_path = os.path.join(CURRENT_DIR, "output")
        if not current_subfolder.startswith(base_output_path):
            print(f"   🛡️ [终极绝对防御] 路径严重偏离正常输出目录！强制重定向！")
            # 强制重建正确的输出目录
            safe_output_dir = os.path.join(CURRENT_DIR, "output", f"{folder_name}_{timestamp}")
            os.makedirs(safe_output_dir, exist_ok=True)
            subfolder_name = f"{batch_index + 1:04d}"
            safe_subfolder = os.path.join(safe_output_dir, subfolder_name)
            os.makedirs(safe_subfolder, exist_ok=True)
            current_subfolder = safe_subfolder
            subfolders[batch_index] = current_subfolder  # 更新缓存
            print(f"   ✅ [已重定向] 新路径: {current_subfolder}")
        
        print(f"\n🔄 进度：第 {i+1}/{total_count} 张 [{prompt_mode}] → 子文件夹 {batch_index + 1:04d}")
        
        # ✨ 生成带前缀的文件名，避免重名冲突
        safe_prefix = folder_name.replace(" ", "_").replace("/", "_")
        filename = f"{target_style}_{i+1:02d}.png"

        # 🛡️ 绝对防御：强行提取正确的子文件夹路径，防止被之前的缓存路径污染！
        # 因为第二张图报错是路径变成了哈希，这里我们强制只要是在输出目录外，就重置。
        if "hf_cache" in os.path.join(current_subfolder, filename) or "snapshots" in os.path.join(current_subfolder, filename):
            # 如果发现目标路径是刚才的哈希缓存目录，直接把最终路径重定向回原来的文件夹！
            print(f"   🛡️ [绝对防御] 检测到路径污染！强制重定向到安全路径...")
            safe_output_dir = os.path.join(CURRENT_DIR, "output", f"{folder_name}_{timestamp}")
            os.makedirs(safe_output_dir, exist_ok=True)
            subfolder_name = f"{batch_index + 1:04d}"
            safe_subfolder = os.path.join(safe_output_dir, subfolder_name)
            os.makedirs(safe_subfolder, exist_ok=True)
            current_subfolder = safe_subfolder
            print(f"   ✅ [已重定向] 新路径: {safe_subfolder}")
            
        # 🟢 [追踪点 1] 调用 generate_style 之前
        print(f"   📁 [追踪 1] 调用 generate_style 前，目标路径: {os.path.join(current_subfolder, filename)}")

        # 🛡️ 修复：决定最终的 strength，并确保传递给 generate_style
        final_strength = DEFAULT_STRENGTH
        # 如果命令行传了 --strength，就用传进来的值
        for arg in sys.argv:
            if arg.startswith('--strength='):
                try:
                    final_strength = float(arg.split('=')[1])
                    break
                except:
                    pass
            elif arg == '--strength' and len(sys.argv) > sys.argv.index(arg) + 1:
                try:
                    final_strength = float(sys.argv[sys.argv.index(arg) + 1])
                    break
                except:
                    pass

        # 🆕 修改：将图片保存到子文件夹
        final_output_path = generate_style(
            pipe, 
            init_image, 
            prompt, 
            os.path.join(current_subfolder, filename),  # 👈 保存到子文件夹
            final_strength, # ✅ 这里必须传 final_strength！
            mode,
            actual_steps,
            target_style,
            # ======= 🆕 新增 IP-Adapter 参数传递 =======
            enable_ip_adapter=enable_ip_adapter,
            ip_adapter_scale=ip_adapter_scale
            # ===========================================            
        )
        saved_img_file = final_output_path

        # 🟢 [追踪点 2] 刚生成完图片 (消除AI痕迹刚刚开始)
        print(f"   ✅ [追踪 2] generate_style 返回的路径 (initial): {final_output_path}")
        print(f"   ✅ [追踪 2] 备份路径 (saved_img_file): {saved_img_file}")
        
        # =================================

        # ====================================================================
        # 🎨 多后端图片鉴赏系统 (复用 gui/backends 全部能力)
        # ====================================================================
        caption = prompt  # 默认使用提示词作为降级选项
        appr_engine = AI_APPRECIATION_ENGINE  # 从 config.py 获取引擎配置

        # 1. 如果设置为 prompt，直接跳过所有模型
        if appr_engine == "prompt":
            print(f"   📝 鉴赏引擎: 仅使用提示词")
            caption = prompt

        # 2. 其他所有引擎：伪装 GUI 界面，复用 backend 类
        else:
            print(f"   🧠 鉴赏引擎: 正在加载 {appr_engine} 后端...")
            try:
                # 创建一个伪装的对象，让 backends 以为自己在 GUI 环境里
                class FakeTab:
                    def __init__(self):
                        self.cancel_interrogate = False
                        self.app = None  # 设为 None，因为我们在命令行下不需要 app

                fake_tab = FakeTab()

                # 🟢 [追踪点 3] 调用后端进行推理前
                print(f"   🔍 [追踪 3] 即将传递给后端进行推理的图片路径: {saved_img_file}")
                
                # 根据配置动态导入后端
                if appr_engine == "tag":
                    from gui.tabs.interrogate.backends.tag import TagBackend
                    backend = TagBackend(fake_tab)
                    # 使用默认的快速标签模式
                    caption = backend.interrogate(saved_img_file, model_name="ViT-Large (准确)", threshold=0.02)
                    
                elif appr_engine == "blip":
                    from gui.tabs.interrogate.backends.blip import BlipBackend
                    backend = BlipBackend(fake_tab)
                    # 使用详细模式
                    caption = backend.interrogate(saved_img_file, model_name="BLIP-large (详细)")
                    
                elif appr_engine == "combined":
                    from gui.tabs.interrogate.backends.combined import CombinedBackend
                    backend = CombinedBackend(fake_tab)
                    # 组合模式：BLIP + CLIP 标签
                    caption = backend.interrogate(saved_img_file, blip_model="BLIP-large (详细)", clip_mode="fast")
                    
                elif appr_engine == "llm":
                    print(f"   🧠 鉴赏引擎: 使用本地目录加载 BLIP + Ollama 润色")
                    caption = prompt
                    
                    # ================= 1. 懒加载 BLIP (只加载一次) =================
                    # 检查当前函数作用域外是否已经加载了模型
                    if 'blip_processor' not in locals() or 'blip_model' not in locals():
                        try:
                            from transformers import BlipProcessor, BlipForConditionalGeneration
                            
                            base_blip_path = r"E:\hf_cache\.cache\hub\models--Salesforce--blip-image-captioning-large"
                            snapshots_path = os.path.join(base_blip_path, "snapshots")
                            
                            if os.path.exists(snapshots_path):
                                subfolders = [f for f in os.listdir(snapshots_path) if os.path.isdir(os.path.join(snapshots_path, f))]
                                if subfolders:
                                    cached_blip_dir = os.path.join(snapshots_path, subfolders[0])
                                    print(f"   📦 首次使用，正在加载 BLIP 模型 ({cached_blip_dir})...")
                                    blip_processor = BlipProcessor.from_pretrained(cached_blip_dir)
                                    blip_model = BlipForConditionalGeneration.from_pretrained(cached_blip_dir)
                                    print(f"   ✅ BLIP 模型加载完成！")
                                else:
                                    raise FileNotFoundError("snapshots 目录为空")
                            else:
                                raise FileNotFoundError(f"找不到 snapshots 目录: {snapshots_path}")
                        except Exception as e:
                            print(f"   ⚠️ 本地 BLIP 加载失败，将使用提示词作为降级。错误: {e}")
                            blip_processor = None
                            blip_model = None

                    # 如果加载成功，则进行推理
                    if blip_processor is not None and blip_model is not None:
                        try:
                            from PIL import Image
                            image = Image.open(saved_img_file).convert('RGB')
                            inputs = blip_processor(image, return_tensors="pt")
                            out = blip_model.generate(**inputs, max_length=80, num_beams=3, repetition_penalty=1.1)
                            blip_caption = blip_processor.decode(out[0], skip_special_tokens=True)
                            print(f"   📝 BLIP 基础描述: {blip_caption[:60]}...")
                        except Exception as e:
                            print(f"   ⚠️ BLIP 推理失败，使用提示词降级。错误: {e}")
                            blip_caption = prompt
                    else:
                        blip_caption = prompt

                    # ================= 2. 将描述提交给本地 Ollama =================
                    if blip_caption and blip_caption != prompt:
                        try:
                            import requests
                            llm_prompt = f"""
请将以下图片描述转换为一段优美、带有艺术鉴赏性的中文赏析（约100字）：
图片描述：{blip_caption}

要求：
1. 包含对人物服装、神态、材质质感的描写。
2. 强调这是一件极具收藏价值的二次元手办/雕像作品。
3. 语言风格：优雅、专业、适合作为社交媒体发帖文案。
"""
                            print(f"   ⏳ 正在请求 Ollama (qwen2.5:1.5b) 润色...")
                            response = requests.post(
                                "http://localhost:11434/api/generate",
                                json={"model": "qwen2.5:1.5b", "prompt": llm_prompt, "stream": False},
                                timeout=45
                            )
                            if response.status_code == 200:
                                caption = response.json().get("response", blip_caption)
                                print(f"   ✅ LLM 润色完成！")
                            else:
                                print(f"   ⚠️ Ollama 返回错误，使用 BLIP 原始描述。")
                                caption = blip_caption
                        except Exception as e:
                            print(f"   ⚠️ Ollama 连接失败。错误: {e}")
                            caption = blip_caption
                    else:
                        caption = blip_caption
                else:
                    print(f"   ⚠️ 未知的引擎配置，使用提示词降级。")
                    caption = prompt

                print(f"   ✅ {appr_engine} 后端推理完成")

            except ImportError as ie:
                print(f"   ❌ 缺少所需依赖，降级为提示词描述。错误: {ie}")
                caption = prompt
            except Exception as e:
                print(f"   ❌ 后端调用失败，降级为提示词描述。错误: {e}")
                caption = prompt
        # ====================================================================

        # ====================================================================
        # 📝 根据后端返回的 caption 生成鉴赏段落
        # ====================================================================
        # 将 AI 返回的图片描述截取前 100 个字，作为鉴赏的核心
        # 🛡️ 修复：如果 caption 是标签堆砌（含 masterpiece 等），则尝试用提示词替代
        if "masterpiece" in caption or "best quality" in caption:
            content_desc = prompt[:100] + "..." if len(prompt) > 100 else prompt
        else:
            content_desc = caption[:100] + "..." if len(caption) > 100 else caption

        # 这段文案就是发帖用的模板，你可以随时修改里面的字眼
        review_paragraph = (
            f"本次 AI 艺术创作描绘了这样一幅画面：“{content_desc}”。\n"
            f"在细腻的笔触和先进的大模型算法加持下，图片不仅呈现出逼真的手办质感，\n"
            f"更通过精准的光影构图，传递出独特的视觉氛围与角色气质。\n"
            f"这是一张兼具技术质感与艺术审美的精致作品。"
        )
        # ====================================================================
        
        # 🛡️ 安全检查：确保 batch_reviews 是个列表
        if 'batch_reviews' not in locals():
            batch_reviews = []

        # 收集当前图片的点评
        batch_reviews.append(f"【第 {len(batch_reviews)+1} 张作品】\n{review_paragraph}")

        # 🆕 判断当前子文件夹是否已经生成完毕（满 5 张，或者总体已经结束）
        is_last_item = (i == total_count - 1)
        is_batch_end = ((i + 1) % BATCH_SIZE == 0)
        
        if is_batch_end or is_last_item:
            # 📄 生成 Word 文档 (严格按 5 张一组)
            if DOCX_AVAILABLE:
                try:
                    import glob
                    doc = Document()
                    
                    # 如果图片少于 BATCH_SIZE 但已经结束了，或者正好凑满 5 张，都会走到这里
                    valid_images = [img for img in glob.glob(os.path.join(current_subfolder, "*.jpg")) if os.path.exists(img)]
                    
                    if valid_images:
                        # 设置标题
                        title = doc.add_heading(f"【{folder_name} 作品合辑】", level=1)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # 添加元数据
                        meta = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        meta.paragraph_format.space_after = Inches(0.2)
                        
                        # 遍历本批次的所有图片和点评
                        for idx, img_path in enumerate(valid_images):
                            review_text = f"【作品 {idx+1}】\n（请在此处微调你的专属评论）"
                            if idx < len(batch_reviews):
                                clean_review = batch_reviews[idx].replace("【", "").replace("】", "").strip()
                                review_text = f"【作品 {idx+1}】\n{clean_review}"

                            # 插入图片 (适应手机屏幕宽度)
                            try:
                                p = doc.add_paragraph()
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = p.add_run()
                                run.add_picture(img_path, width=Inches(5.5))
                            except Exception as e:
                                print(f"      ⚠️ Word 插入图片失败: {e}")

                            # 插入图片下方的鉴赏文字
                            review_p = doc.add_paragraph(review_text)
                            review_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            review_p.paragraph_format.space_before = Inches(0.1)
                            review_p.paragraph_format.space_after = Inches(0.3)

                        # 保存 Word 文档
                        docx_file = os.path.join(current_subfolder, "公众号草稿.docx")
                        doc.save(docx_file)
                        print(f"      📄 已生成可直接导入公众号的 Word 文档：{os.path.basename(docx_file)}")
                    else:
                        print(f"      ⚠️ 本组未发现有效 JPG 图片，跳过 Word 生成。")
                except Exception as e:
                    print(f"      ⚠️ Word 文档生成失败：{e}")
            else:
                print(f"      ℹ️ 跳过 Word 生成 (python-docx 未安装)")

            # ========== 📝 保留 Txt 备份 ==========
            try:
                summary_file = os.path.join(current_subfolder, "点评.txt")
                with open(summary_file, "w", encoding="utf-8") as f:
                    f.write(f"【{folder_name} AI 作品鉴赏合辑】\n")
                    f.write(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                    f.write(f"本集共收录 {len(batch_reviews)} 件 AI 视觉创作：\n\n")
                    
                    for review in batch_reviews:
                        f.write(f"{review}\n\n")
                        
                    #f.write(f"—— 由 AI 视觉鉴赏系统自动书写 ——\n")
                
                print(f"      📝 已生成备份 Txt 文档：{os.path.basename(summary_file)}")
            except Exception as e:
                print(f"      ⚠️ Txt 备份文档写入失败：{e}")

            # 重置缓存，准备下一个文件夹
            batch_reviews = []

    print(f"\n✅ 全部完成！共 {total_count} 张图片，保存在: {output_root}")
    print(f"📁 图片已按每 {BATCH_SIZE} 张分到 {total_batches} 个子文件夹中")

if __name__ == "__main__":
    main()